from docx import Document
import os
import json

def fill_resume_template(data: dict, template_path='templates/resume_template.docx', output_folder='output') -> str:
    doc = Document(template_path)

    for paragraph in doc.paragraphs:
        inline_replacement(paragraph, data)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    inline_replacement(paragraph, data)

    os.makedirs(output_folder, exist_ok=True)
    file_name = f"resume_{data.get('full_name', 'user').replace(' ', '_')}.docx"
    output_path = os.path.join(output_folder, file_name)
    doc.save(output_path)
    print(f"✅ Resume saved to: {output_path}")
    return output_path

def inline_replacement(paragraph, data):
    full_text = ''.join(run.text for run in paragraph.runs)
    

    # Handle special case for contact URLs
    if "{{urls}}" in full_text:
        contact_parts = []
        
        # Add email if available and not None/empty
        email = data.get('email_address') or data.get('email')
        if email and str(email).strip() and str(email).lower() != 'none':
            contact_parts.append(str(email).strip())
        
        # Add phone if available and not None/empty
        phone = data.get('phone_number') or data.get('phone')
        if phone and str(phone).strip() and str(phone).lower() != 'none':
            contact_parts.append(str(phone).strip())
        
        # Add LinkedIn if available and not None/empty
        linkedin = data.get('linkedin_url') or data.get('linkedin')
        if linkedin and str(linkedin).strip() and str(linkedin).lower() != 'none':
            contact_parts.append(str(linkedin).strip())
        
        # Add GitHub if available and not None/empty
        github = data.get('github_url') or data.get('github')
        if github and str(github).strip() and str(github).lower() != 'none':
            contact_parts.append(str(github).strip())
        
        # Join with " | " separator
        urls_value = " | ".join(contact_parts)
        full_text = full_text.replace("{{urls}}", urls_value)
        
        # Clear and rebuild paragraph
        for i in range(len(paragraph.runs) - 1, -1, -1):
            paragraph._element.remove(paragraph.runs[i]._element)
        paragraph.add_run(full_text)
        return

    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in full_text:
            if isinstance(value, list):
                # Handle bullet lists with proper Word formatting
                handle_bullet_list(paragraph, value, placeholder, full_text)
                return  # Exit early since we've handled this paragraph
            elif isinstance(value, dict):
                # Convert dicts (e.g., one project) into formatted string
                formatted = ''
                if 'title' in value and 'description' in value:
                    formatted = f"{value['title']}: {value['description']}"
                else:
                    formatted = json.dumps(value, indent=2)
                value = formatted
            else:
                value = str(value)

            # Replace and rebuild paragraph for non-list items
            full_text = full_text.replace(placeholder, value)

            for i in range(len(paragraph.runs) - 1, -1, -1):
                paragraph._element.remove(paragraph.runs[i]._element)
            paragraph.add_run(full_text)


def handle_bullet_list(paragraph, items, placeholder, full_text):
    """Handle bullet lists by creating properly formatted paragraphs"""
    from docx.shared import Inches
    
    # Get the parent element (document or table cell)
    parent = paragraph._element.getparent()
    
    # Find the position of the current paragraph
    para_index = list(parent).index(paragraph._element)
    
    # Process each item in the list
    for i, item in enumerate(items):
        if i == 0:
            # Use the existing paragraph for the first item
            current_para = paragraph
        else:
            # Create new paragraph for subsequent items
            new_para_element = parent.makeelement(paragraph._element.tag)
            parent.insert(para_index + i, new_para_element)
            # Create paragraph object from element
            from docx.text.paragraph import Paragraph
            current_para = Paragraph(new_para_element, paragraph._parent)
        
        # Clear existing runs
        for j in range(len(current_para.runs) - 1, -1, -1):
            current_para._element.remove(current_para.runs[j]._element)
        
        # Format the bullet item
        if isinstance(item, str):
            bullet_text = f"• {item}"
        elif isinstance(item, dict):
            line = item.get("title", "")
            if "provider" in item:
                line += f", {item['provider']}"
            if "date" in item:
                line += f" ({item['date']})"
            
            bullet_text = f"• {line}"
            
            # Add description as a separate line if it exists
            if "description" in item:
                bullet_text += f"\n{item['description']}"
        
        # Add the formatted text to the paragraph
        current_para.add_run(bullet_text)
        
        # Apply hanging indent formatting
        try:
            current_para.paragraph_format.left_indent = Inches(0.25)
            current_para.paragraph_format.first_line_indent = Inches(-0.25)
        except:
            pass


def load_json(json_path):
    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)

# if __name__ == "__main__":
#     data = load_json("output/resume_Manjunath.json")  # update path if needed
#     data["full_name"] = "Manjunath"  # Ensure it's there for file naming
#     fill_resume_template(data)
